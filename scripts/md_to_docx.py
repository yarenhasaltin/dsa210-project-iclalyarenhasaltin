#!/usr/bin/env python3
"""Convert FINAL_REPORT.md to FINAL_REPORT.docx with cleaned text."""

import re
from pathlib import Path
from docx import Document

INPUT_PATH = Path(__file__).resolve().parent.parent / "FINAL_REPORT.md"
OUTPUT_PATH = Path(__file__).resolve().parent.parent / "FINAL_REPORT.docx"

EMOJI_PATTERN = re.compile(
    r"[\U0001F600-\U0001F64F]|[\U0001F300-\U0001F5FF]|[\U0001F680-\U0001F6FF]|"
    r"[\U0001F700-\U0001F77F]|[\U0001F780-\U0001F7FF]|[\U0001F800-\U0001F8FF]|"
    r"[\U0001F900-\U0001F9FF]|[\U0001FA00-\U0001FA6F]|[\U0001FA70-\U0001FAFF]|"
    r"[\u2600-\u26FF]|[\u2700-\u27BF]"
)

HEADER_PATTERN = re.compile(r"^(#{1,6})\s*(.*)$")


def clean_text(text: str) -> str:
    text = EMOJI_PATTERN.sub("", text)
    text = text.replace("✅", "")
    text = text.replace("⚠️", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def add_markdown_paragraph(doc: Document, line: str) -> None:
    match = HEADER_PATTERN.match(line)
    if match:
        level = len(match.group(1))
        text = clean_text(match.group(2))
        if not text:
            return
        if level == 1:
            doc.add_heading(text, level=1)
        elif level == 2:
            doc.add_heading(text, level=2)
        elif level == 3:
            doc.add_heading(text, level=3)
        else:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
    elif line.startswith("- ") or line.startswith("* "):
        text = clean_text(line[2:])
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(text)
    elif line.startswith("1. ") or re.match(r"^\d+\. ", line):
        text = clean_text(re.sub(r"^\d+\. ", "", line))
        para = doc.add_paragraph(style="List Number")
        para.add_run(text)
    elif line.strip() == "" or line.startswith("---"):
        return
    else:
        text = clean_text(line)
        doc.add_paragraph(text)


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    doc = Document()
    with input_path.open("r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip("\n")
        add_markdown_paragraph(doc, line)

    doc.save(output_path)
    print(f"Saved DOCX to: {output_path}")


if __name__ == "__main__":
    convert_markdown_to_docx(INPUT_PATH, OUTPUT_PATH)
